"""
Conversor Markdown -> DOCX generalizado (encabezados, tablas, listas, citas,
negritas e IMÁGENES ![alt](ruta)). Uso:
    python scripts/md_to_docx_general.py <entrada.md> [salida.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

IN_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('docs/fase3_modelado.md')
OUT_PATH = Path(sys.argv[2]) if len(sys.argv) > 2 else IN_PATH.with_suffix('.docx')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15
for section in doc.sections:
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

content = IN_PATH.read_text(encoding='utf-8')
lines = content.split('\n')
IMG_RE = re.compile(r'!\[(.*?)\]\((.*?)\)')


def clean_md(text):
    text = IMG_RE.sub('', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text.strip()


def add_rich_paragraph(p, text):
    """Añade runs con **negrita** y `código` a un párrafo."""
    text = re.sub(r'`(.*?)`', r'\1', text)
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)


def add_image(path_str, alt):
    img = (IN_PATH.parent / path_str).resolve()
    if not img.exists():
        doc.add_paragraph(f'[imagen no encontrada: {path_str}]')
        return
    doc.add_picture(str(img), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.runs[0]; r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(110, 110, 110)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]; cell.text = ''
        run = cell.paragraphs[0].add_run(clean_md(h))
        run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._element.get_or_add_tcPr()
        tcPr.append(tcPr.makeelement(qn('w:shd'),
                    {qn('w:fill'): '2E4057', qn('w:val'): 'clear', qn('w:color'): 'auto'}))
    for i, row in enumerate(rows):
        for j in range(min(len(row), len(headers))):
            cell = table.rows[i + 1].cells[j]; cell.text = ''
            p = cell.paragraphs[0]
            add_rich_paragraph(p, row[j])
            for run in p.runs:
                run.font.size = Pt(10)
            if i % 2 == 1:
                tcPr = cell._element.get_or_add_tcPr()
                tcPr.append(tcPr.makeelement(qn('w:shd'),
                            {qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear', qn('w:color'): 'auto'}))
    doc.add_paragraph()


def parse_table(lines, start):
    raw0 = lines[start].strip().strip('|')
    headers = [c.strip() for c in raw0.split('|')]
    i = start + 1
    if i < len(lines) and re.match(r'^[\s|:-]+$', lines[i]):
        i += 1
    rows = []
    while i < len(lines) and lines[i].strip() and '|' in lines[i]:
        raw = lines[i].strip()
        raw = raw[1:] if raw.startswith('|') else raw
        raw = raw[:-1] if raw.endswith('|') else raw
        rows.append([c.strip() for c in raw.split('|')])
        i += 1
    return headers, rows, i


i = 0
while i < len(lines):
    stripped = lines[i].strip()
    if not stripped or stripped == '---':
        i += 1; continue

    m = IMG_RE.search(stripped)
    if m and stripped.startswith('!['):
        add_image(m.group(2), m.group(1)); i += 1; continue

    if stripped.startswith('# ') and not stripped.startswith('## '):
        h = doc.add_heading(clean_md(stripped[2:]), level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs: r.font.color.rgb = RGBColor(30, 50, 80)
        i += 1; continue
    if stripped.startswith('## '):
        h = doc.add_heading(clean_md(stripped[3:]), level=2)
        for r in h.runs: r.font.color.rgb = RGBColor(46, 64, 87)
        i += 1; continue
    if stripped.startswith('### '):
        doc.add_heading(clean_md(stripped[4:]), level=3); i += 1; continue
    if stripped.startswith('#### '):
        doc.add_heading(clean_md(stripped[5:]), level=4); i += 1; continue

    if '|' in stripped and not stripped.startswith('>'):
        headers, rows, end = parse_table(lines, i)
        if headers and rows:
            add_table(headers, rows)
        i = end; continue

    if stripped.startswith('- ') or stripped.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        add_rich_paragraph(p, stripped[2:]); i += 1; continue

    if stripped.startswith('> '):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
        add_rich_paragraph(p, stripped[2:])
        for r in p.runs:
            r.italic = True; r.font.color.rgb = RGBColor(90, 90, 90)
        i += 1; continue

    p = doc.add_paragraph()
    add_rich_paragraph(p, stripped)
    i += 1

doc.save(str(OUT_PATH))
print(f'Documento guardado: {OUT_PATH}')
