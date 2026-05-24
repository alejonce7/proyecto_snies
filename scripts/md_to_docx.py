"""
Convierte fase1_stakeholders_preguntas.md a .docx 
Adaptado al formato de la Fase 1
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# Configurar estilos
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Leer el markdown
with open('docs/fase1_stakeholders_preguntas.md', 'r', encoding='utf-8') as f:
    content = f.read()
    lines = content.split('\n')

def clean_md(text):
    """Remueve markdown formatting"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = text.strip()
    return text

def add_rich_paragraph(doc, text):
    """Agrega párrafo con bold donde haya **texto**"""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_table(doc, headers, rows):
    """Agrega tabla formateada"""
    if not headers:
        return
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers con fondo azul oscuro
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Calibri'
        # Fondo azul
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn('w:shd'), {
            qn('w:fill'): '2E4057',
            qn('w:val'): 'clear',
            qn('w:color'): 'auto'
        })
        tcPr.append(shading)
    
    # Filas con alternado gris claro
    for i, row in enumerate(rows):
        for j in range(min(len(row), ncols)):
            cell = table.rows[i+1].cells[j]
            # Check marks y contenido
            val = row[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            
            # Centrar columnas con checkmarks
            if val.strip() in ['✔', '✔️', '✅', '']:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fondo alternado
            if i % 2 == 1:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shading = tcPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F2F2',
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto'
                })
                tcPr.append(shading)
    
    doc.add_paragraph()

def parse_table_block(lines, start):
    """Parsea tabla markdown"""
    headers = []
    rows = []
    i = start
    
    if '|' in lines[i]:
        headers = [clean_md(c.strip()) for c in lines[i].strip().strip('|').split('|') if c.strip() != '']
        i += 1
    
    if i < len(lines) and '---' in lines[i]:
        i += 1
    
    while i < len(lines) and lines[i].strip() and '|' in lines[i]:
        cells = [clean_md(c.strip()) for c in lines[i].strip().strip('|').split('|') if c.strip() != '' or lines[i].count('|') > 2]
        # Handle empty cells properly
        raw = lines[i].strip()
        if raw.startswith('|'):
            raw = raw[1:]
        if raw.endswith('|'):
            raw = raw[:-1]
        cells = [clean_md(c.strip()) for c in raw.split('|')]
        rows.append(cells)
        i += 1
    
    return headers, rows, i

# Procesar
i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    if not stripped:
        i += 1
        continue
    
    if stripped == '---':
        i += 1
        continue
    
    # H1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        title = clean_md(stripped[2:])
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 50, 80)
        i += 1
        continue
    
    # H2
    if stripped.startswith('## '):
        title = clean_md(stripped[3:])
        h = doc.add_heading(title, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(46, 64, 87)
        i += 1
        continue
    
    # H3
    if stripped.startswith('### '):
        title = clean_md(stripped[4:])
        doc.add_heading(title, level=3)
        i += 1
        continue
    
    # H4
    if stripped.startswith('#### '):
        title = clean_md(stripped[5:])
        doc.add_heading(title, level=4)
        i += 1
        continue
    
    # Tabla
    if '|' in stripped and not stripped.startswith('>'):
        headers, rows, end = parse_table_block(lines, i)
        if headers and rows:
            add_table(doc, headers, rows)
        i = end
        continue
    
    # Lista con bullet
    if stripped.startswith('- '):
        text = stripped[2:]
        if '**' in text:
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(text, style='List Bullet')
        i += 1
        continue
    
    # Blockquote
    if stripped.startswith('> '):
        text = clean_md(stripped[2:])
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        i += 1
        continue
    
    # Texto normal con formato
    if '**' in stripped:
        add_rich_paragraph(doc, stripped)
    else:
        doc.add_paragraph(stripped)
    
    i += 1

# Guardar
output_path = 'docs/fase1_stakeholders_preguntas.docx'
doc.save(output_path)
print(f"Documento guardado: {output_path}")
